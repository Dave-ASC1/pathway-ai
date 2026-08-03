#!/usr/bin/env python3
"""
Builds Pathway_AI_Project_Management_Charts_Updated.docx
Mirrors the original charts document, updated to July 6, 2026 status,
embedding the three regenerated charts (gantt/cpa/pert _updated.png).

Usage:  pip install python-docx ; python3 build-charts-doc.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

NAVY = RGBColor(0x04, 0x13, 0x36)
BLUE = RGBColor(0x3b, 0x82, 0xf6)
INK  = RGBColor(0x0f, 0x17, 0x2a)
MUTED = RGBColor(0x64, 0x74, 0x8b)
FONT = "Arial"
HERE = os.path.dirname(os.path.abspath(__file__))

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(11)
normal.font.color.rgb = INK

def run(p, text, size=11, bold=False, color=INK, italic=False):
    r = p.add_run(text); r.font.name = FONT; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color
    return r

def para(text="", size=11, bold=False, color=INK, align=None, after=8, italic=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    if align: p.alignment = align
    if text: run(p, text, size, bold, color, italic)
    return p

def heading(text, size=16):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run(p, text, size, True, NAVY)
    return p

def figure(path, caption, width=6.6):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.add_run().add_picture(path, width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(12)
    run(c, caption, 9, False, MUTED, italic=True)

def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hex_fill)
    tcPr.append(sh)

# ---- Title block ----
para("PATHWAY AI", 26, True, NAVY, WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("Project Management Charts (Updated)", 16, False, BLUE, WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("IST 440W Capstone  |  Penn State University", 11, False, MUTED, WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("David Ademoye  |  May 20 – August 12, 2026  |  Status as of July 6, 2026",
     11, False, MUTED, WD_ALIGN_PARAGRAPH.CENTER, after=14)

# ---- 1. Summary ----
heading("1. Project Management Summary")
para("This document contains the updated project management artifacts for Pathway AI, "
     "reflecting project status as of July 6, 2026. Since the June 30 baseline, the project "
     "has advanced well ahead of the original plan. Vercel deployment is complete with a "
     "confirmed public URL, the Claude Opus 4.8 API is integrated, and all four career "
     "readiness modules (Resume Checker, Career Path Explorer, Skill Gap Roadmap, and Mock "
     "Interview Coach) are built and AI-powered.")
para("The remaining critical path is short: Testing and QA, followed by final documentation "
     "and submission, closing with the capstone presentation on August 12. Clerk authentication "
     "and the NeonDB database were intentionally moved to a parallel, non-critical track. Because "
     "the application uses browser-based storage, these items carry available slack and can be "
     "completed as stretch goals without affecting the submission date.")

# ---- 2. Gantt ----
heading("2. Gantt Chart")
para("The Gantt chart plots every task across the May 20 to August 12 timeline. Green bars are "
     "completed work, amber bars are currently in progress, and blue bars are planned stretch "
     "items. The red dashed line marks today, July 6, 2026.")
figure(os.path.join(HERE, "gantt_updated.png"),
       "Figure 1. Pathway AI Gantt Chart (Updated) — May 20 to August 12, 2026")

# ---- 3. CPA ----
heading("3. Critical Path Analysis")
para("The critical path is the sequence of dependent tasks that controls the final delivery date. "
     "With the core build complete, the critical path now runs only through Testing and QA and "
     "Final Documentation and Submission. Authentication and the database have moved off the "
     "critical path into a parallel track with slack.")
figure(os.path.join(HERE, "cpa_updated.png"),
       "Figure 2. Pathway AI Critical Path (Updated) — core build complete")
para("Updated Critical Path Sequence", 12, True, NAVY, after=4)
for line in [
    "A  Requirements & Brief (5d)  — Completed",
    "C  Design & Branding (8d)  — Completed",
    "D  Dashboard Route (5d)  — Completed",
    "E  Resume Checker MVP (8d)  — Completed",
    "F  Vercel Deployment (4d)  — Completed",
    "G  Claude API Integration (12d)  — Completed",
    "I / J / K  Four AI Modules Built (approx. 20d)  — Completed",
    "Testing & QA (10d)  — In Progress",
    "L  Final Documentation & Submission (8d)  — In Progress",
]:
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
    run(p, line, 11)
para("The completed and in-progress critical path totals roughly 80 expected days within the "
     "84-day window. The core development is finished ahead of schedule, leaving comfortable "
     "float for testing, documentation, and the optional authentication and database work.",
     11, after=8)

# ---- 4. PERT table ----
heading("4. PERT Estimate Table")
para("Each task is estimated with three scenarios: Optimistic (O), Most Likely (M), and "
     "Pessimistic (P). PERT expected time is te = (O + 4M + P) / 6. Status is updated to July 6, 2026.")
rows = [
    ("A","Requirements & Project Brief","3","5","7","5.0","Completed"),
    ("B","Literature Review & Research","5","7","10","7.2","Completed"),
    ("C","Design & Branding","5","8","12","8.2","Completed"),
    ("D","Dashboard Route","3","5","8","5.2","Completed"),
    ("E","Resume Checker MVP","5","8","12","8.2","Completed"),
    ("F","Vercel Deployment","2","4","7","4.2","Completed"),
    ("G","Claude API Integration","7","12","18","12.2","Completed"),
    ("H","Clerk Auth + NeonDB","7","10","15","10.3","Planned"),
    ("I","Career Path Explorer","7","12","18","12.2","Completed"),
    ("J","Skill Gap Roadmap","7","12","18","12.2","Completed"),
    ("K","Mock Interview Coach","10","14","21","14.5","Completed"),
    ("L","Final Documentation & Submission","5","8","12","8.2","In Progress"),
]
headers = ["ID","Task","O","M","P","te","Status"]
table = doc.add_table(rows=1, cols=7)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
hdr = table.rows[0].cells
for i, h in enumerate(headers):
    shade(hdr[i], "041336")
    pph = hdr[i].paragraphs[0]; pph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(pph, h, 10, True, RGBColor(0xff,0xff,0xff))
status_color = {"Completed": RGBColor(0x15,0x80,0x3d),
                "In Progress": RGBColor(0xb4,0x5c,0x00),
                "Planned": RGBColor(0x1d,0x4e,0xd8)}
for ri, r in enumerate(rows):
    cells = table.add_row().cells
    for ci, val in enumerate(r):
        if ri % 2 == 1: shade(cells[ci], "f7f9fc")
        pp = cells[ci].paragraphs[0]
        if ci in (0,2,3,4,5): pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        col = status_color.get(val, INK) if ci == 6 else INK
        run(pp, val, 10, ci==6, col)
para("O = Optimistic days  |  M = Most Likely days  |  P = Pessimistic days  |  te = PERT expected time",
     9, False, MUTED, after=10)

# ---- 5. PERT chart ----
heading("5. PERT Chart")
para("The PERT chart maps all tasks as nodes in a dependency network. Each node shows the task "
     "ID, name, and expected duration (te). Arrows represent dependencies: a task cannot begin "
     "until all upstream tasks are complete. As of July 6, every build node is complete; only "
     "node H (Auth and Database) remains planned, and it sits on a parallel track.")
figure(os.path.join(HERE, "pert_updated.png"),
       "Figure 3. Pathway AI PERT Chart (Updated) — May 20 to August 12, 2026")
para("The two development tracks (frontend and AI integration, A → C → E → G → I → L; and "
     "infrastructure, B → D → F → H → K → L) both converge at final submission. With the AI "
     "track fully complete, delivery risk now rests almost entirely on documentation and testing, "
     "both of which are underway with slack to spare.")

out = os.path.join(HERE, "Pathway_AI_Project_Management_Charts_Updated.docx")
doc.save(out)
print("wrote", out)
