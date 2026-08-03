#!/usr/bin/env python3
"""
Builds Pathway_AI_Development_Prompts_Log.docx
A curated log of the development/session prompts David used with Claude to
design and build Pathway AI. Each entry pairs a short purpose with the
verbatim prompt. Matches the brand styling of the other capstone documents.

Usage: pip install python-docx ; python3 build-prompts-doc.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

NAVY  = RGBColor(0x04, 0x13, 0x36)
BLUE  = RGBColor(0x3b, 0x82, 0xf6)
INK   = RGBColor(0x0f, 0x17, 0x2a)
MUTED = RGBColor(0x64, 0x74, 0x8b)
FONT  = "Arial"
HERE  = os.path.dirname(os.path.abspath(__file__))

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
    s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)
normal = doc.styles["Normal"]
normal.font.name = FONT; normal.font.size = Pt(11); normal.font.color.rgb = INK

def run(p, text, size=11, bold=False, color=INK, italic=False):
    r = p.add_run(text); r.font.name = FONT; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color
    return r

def para(text="", size=11, bold=False, color=INK, align=None, after=8, italic=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    if align: p.alignment = align
    if text: run(p, text, size, bold, color, italic)
    return p

def heading(text, size=15, before=14):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(5); run(p, text, size, True, NAVY)
    return p

def shade(paragraph, hex_fill):
    pPr = paragraph._p.get_or_add_pPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hex_fill)
    pPr.append(sh)
    # a little inner spacing via border spacing
    for edge in ("left", "right"):
        pass

def prompt_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.12); p.paragraph_format.right_indent = Inches(0.12)
    shade(p, "f3f6fb")
    run(p, text, 10.5, False, INK)

# ---- Title ----
para("PATHWAY AI", 26, True, NAVY, WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("AI Development Prompt Log", 16, False, BLUE, WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("IST 440W Capstone  |  Penn State University", 11, False, MUTED, WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("David Ademoye  |  July 6, 2026", 11, False, MUTED, WD_ALIGN_PARAGRAPH.CENTER, after=14)

para("This log records the development prompts used with Claude to design, redesign, and build "
     "Pathway AI during the enhancement phase. These are the author's conversational prompts to "
     "the AI assistant, distinct from the five application prompts the product itself sends to the "
     "Claude API (documented in the Enhancement and Implementation Report). Each entry lists the "
     "intent behind the prompt followed by the prompt as written.", after=12)

prompts = [
    ("1. Project Context and Guardrails",
     "Establish full project context at the start of a session, point the assistant to the "
     "authoritative handoff and architecture docs, require it to verify current state before acting, "
     "and lock in the brand rules (always show Pathway AI, never Claude or Anthropic; no em or en dashes).",
     "I'm continuing work on Pathway AI, a free student career-readiness web app (Next.js 16 + Claude API). "
     "Four AI-powered tools, all live in production: resume checker, career path explorer, skill gap roadmap, "
     "mock interview coach. The entry experience is an interactive \"journey board\" (a game-style switchback "
     "road with an auto-driving car and fireworks) that serves as both the landing page and the dashboard. "
     "Project path: /Users/kingdavid/pathway-ai GitHub: https://github.com/Dave-ASC1/pathway-ai (branch main, "
     "auto-deploys to Vercel) Live URL: https://pathway-aiapp.vercel.app Before doing anything else, read these "
     "two files in the repo for full context: 1. docs/PathwayAI_New_LLM_Project_Handoff.md - full project "
     "history, decisions, and current state (see story item #12 for the journey board). 2. AGENTS.md - "
     "technical architecture reference (code patterns, brand rules, the Journey Board component, known gotchas). "
     "(CLAUDE.md is auto-loaded but stale - trust the handoff and live code over it.) Then verify the state is "
     "what those docs describe rather than assuming: run git log --oneline -10 (HEAD should be at or after "
     "43a79db), confirm local matches origin/main, and spot-check the live URL (curl the landing page for "
     "\"Reach career-ready in 4 steps\") before treating anything as current. Two rules to hold from the start: "
     "never show \"Claude\" or \"Anthropic\" in the UI (always \"Pathway AI\") and never use em dashes or en "
     "dashes in copy or in AI-generated output. Once you've confirmed the current state, ask me what I want to "
     "work on next."),

    ("2. Landing Page UX Review and Headline",
     "Request a professional UX/UI critique of the landing page and other pages, aimed at making them "
     "more welcoming and improving retention, and flag that the current headline reads awkwardly.",
     "The first landing page, You are a professional UX/UI designer and you have been contracted by my "
     "company to make the landing page and the other pages welcoming, user friendly and retain their "
     "patronage, what do you think can be done to make it better? the text \"Reach career-ready in 4 steps\" "
     "doesn't sound smooth."),

    ("3. Headline Fix and Journey Redesign Concept",
     "Fix the headline and address broader layout and content gaps, then pitch the core redesign idea: a "
     "larger scrolling road where each stop is its own focus, with a car that drives down between stops "
     "and a fireworks finish-line banner on completion.",
     "Fix the headline and take a pass at the bigger layout/content gaps. also, I was thinking about another "
     "thing. instead of the way the entire roadmap is comprised onto one page, how about we make it a larger "
     "road with each stop focusing on a particular stop, I don't know if the car should be heading down or up "
     "(I'd say down since the page scrolls down) so when the user lands on the page, the first thing they would "
     "see is the start your journey or start here button with a 2 car driving down to the next goal the moment "
     "the user clicks on button, then the car starts its journey down and stops at the next point then again "
     "until the finish line, the finish line should change to a finish line banner with stars on either side of "
     "the pole and fireworks erupting once the user gets through the entire journey."),

    ("4. Alternate Layout Exploration (Left to Right)",
     "Explore a horizontal variant of the concept where the car drives left to right and each completed step "
     "fuels the car toward a finish line on the far right.",
     "I think we should do a overhaul of how the entire site is viewed, do you think a left to right view would "
     "be possible, where the car would be driving left to right, and each page you click and step you complete "
     "would get the car fuel to go to the next phase moving further right until it gets all the way to the "
     "finish line."),

    ("5. Course Correction on the Road Shape",
     "Reject the horizontal direction and specify the intended visual: a road that curves down from left to "
     "center then downward to the next stop, continuing after each phase is completed.",
     "Nope, not the vision, the road should be trending down like a curve going from left to center of the "
     "screen then downwards to the next roadmap, then after completion of that phase the road continues."),

    ("6. Results Visualization Decision",
     "Decide how to present resume section scores on the full-width dashboard: whether to replace the section "
     "cards with a single donut chart (slices per section, colored by score band) or keep both.",
     "Pie chart applies to both, Dashboard (full width), and do you think it is visually better to replace the "
     "section cards with a single donut pie or can we have both, where we have the donut pie chart where each "
     "slice represents a section, colored by score band and score identity."),

    ("7. Curate Final Deliverables",
     "Compile the full set of submission artifacts: the curated prompt log, the code, the generated document, "
     "and the updated Gantt chart, Critical Path Analysis, and PERT chart.",
     "Create the curated list of all my prompts from this session, my codes and the generated document as well "
     "as an updated Gantt Chart, Critical Path Analysis, PERT chart."),
]

for title, purpose, text in prompts:
    heading(title)
    pp = doc.add_paragraph(); pp.paragraph_format.space_after = Pt(4)
    run(pp, "Purpose: ", 11, True, NAVY); run(pp, purpose, 11, False, INK)
    lp = doc.add_paragraph(); lp.paragraph_format.space_after = Pt(2)
    run(lp, "Prompt as written:", 10.5, True, MUTED)
    prompt_block(text)

out = os.path.join(HERE, "Pathway_AI_Development_Prompts_Log.docx")
doc.save(out)
print("wrote", out)
