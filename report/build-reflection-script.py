#!/usr/bin/env python3
"""
Builds Pathway_AI_Video_Reflection_Script.docx
A spoken-word, teleprompter-friendly self-reflection script for David's
video reflection. Large readable body text, section headers with time cues,
natural first-person cadence, no em or en dashes.

Usage: pip install python-docx ; python3 build-reflection-script.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

NAVY  = RGBColor(0x04, 0x13, 0x36)
BLUE  = RGBColor(0x3b, 0x82, 0xf6)
INK   = RGBColor(0x0f, 0x17, 0x2a)
MUTED = RGBColor(0x64, 0x74, 0x8b)
FONT  = "Arial"
HERE  = os.path.dirname(os.path.abspath(__file__))

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
    s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)
normal = doc.styles["Normal"]
normal.font.name = FONT; normal.font.size = Pt(13); normal.font.color.rgb = INK

def run(p, text, size=13, bold=False, color=INK, italic=False):
    r = p.add_run(text); r.font.name = FONT; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color
    return r

def para(text="", size=13, bold=False, color=INK, align=None, after=10, line=1.4, italic=False):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_after = Pt(after); pf.line_spacing = line
    if align: p.alignment = align
    if text: run(p, text, size, bold, color, italic)
    return p

def cue(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run(p, text, 12, True, BLUE)
    return p

# ---- Header ----
para("Pathway AI  -  Video Reflection Script", 20, True, NAVY, WD_ALIGN_PARAGRAPH.CENTER, after=4, line=1.1)
para("David Ademoye  |  IST 440W Capstone  |  Penn State University",
     11, False, MUTED, WD_ALIGN_PARAGRAPH.CENTER, after=10, line=1.1)

tip = doc.add_paragraph(); tip.paragraph_format.space_after = Pt(4)
run(tip, "How to use this: ", 11, True, NAVY)
run(tip, "Read at a calm, natural pace (about 130 to 150 words per minute). "
         "Full read time is roughly 7 to 8 minutes. The blue lines are section cues and "
         "timing guides, not meant to be read aloud. Pause briefly wherever a paragraph ends.",
    11, False, MUTED)
doc.add_paragraph()

# ---- Script body ----
cue("OPENING  (about 0:00 - 0:30)")
para("Hi, my name is David Ademoye, and this is my capstone reflection on Pathway AI. "
     "Pathway AI is a free career readiness platform for college students. It brings four tools "
     "together in one place: an ATS resume checker, a career path explorer, a skill gap roadmap, "
     "and a mock interview coach. Tools like these usually live in separate paid products. My goal "
     "was to combine them, make them genuinely useful, and give them away for free. Over the next "
     "few minutes I want to walk you through what I built, what I learned, and where I struggled.")

cue("THE TRANSFORMATION  (about 0:30 - 2:00)")
para("When I started, the resume checker already worked, but it worked in a shallow way. It matched "
     "keywords. If your resume was missing the word Python, it told you so. That was technically "
     "correct, but it felt hollow. It could not tell you why Python mattered for the role you wanted, "
     "or how to show that skill honestly. The other three modules were just placeholder cards. The app "
     "looked finished, but it only delivered real value in one place.")
para("The turning point was integrating Claude. I moved every module onto the Claude API, and that "
     "single decision changed the project from a nice looking wrapper around simple logic into a tool "
     "that actually understands context. The resume checker now reads your resume against a real job "
     "description and scores it the way a recruiter might. The career explorer reasons about your major "
     "and your interests. The skill gap roadmap tells you what to learn next and in what order. The "
     "interview coach asks you real questions and grades your answers. That is the difference between "
     "counting words and giving guidance.")

cue("WHAT I LEARNED  (about 2:00 - 4:15)")
para("The biggest thing I learned is that prompt engineering is a craft, not a magic trick. I expected "
     "Claude to just work once I handed it a prompt. Instead, my early prompts were too loose. The model "
     "would wrap its answers in extra formatting, or return two career paths when I asked for three. The "
     "fix was not clever. It was iteration. I would write a prompt, test it ten times with very different "
     "inputs, find where it broke, and tighten it. Adding one clear instruction, like return only valid "
     "data with exactly three paths, took my reliability from good to almost perfect. That taught me to "
     "think about how the model reads instructions, not just what I want it to do.")
para("The second lesson was about reliability. Real systems need a backup plan. If the AI is slow, rate "
     "limited, or unavailable, the app cannot just crash. So I kept the original keyword logic as a local "
     "fallback. If the AI call fails, the app quietly falls back to the deterministic version and keeps "
     "working. I also added rate limiting on each feature so the service stays fair and protected. It is "
     "not glamorous work, but it is the difference between a demo and something people can actually rely on.")
para("The third lesson was about integration. The whole point of Pathway AI is that the tools talk to "
     "each other. Your resume and your skills carry from one module to the next, so you never have to re "
     "enter the same information. That small detail is what makes it feel like one journey instead of four "
     "separate apps glued together.")

cue("RESPONDING TO FEEDBACK  (about 4:15 - 5:15)")
para("I also want to be honest about feedback. After a review, my professor pointed out that the "
     "experience needed to be more captivating for students. My first version was functional but plain. "
     "It looked like a form. So I redesigned the entry experience around a journey metaphor. The landing "
     "page became an interactive road with four stops, a car that drives along the path, and a finish line "
     "the user reaches when they complete the journey. I reworked the results view too, with a clear match "
     "score and a visual breakdown of each resume section. That feedback pushed the project from something "
     "that worked to something that invites you in, and I am grateful for it.")

cue("THE HARDEST PART  (about 5:15 - 6:15)")
para("The hardest part was not the code. It was deciding what not to build. With six weeks and one person, "
     "I had to constantly choose between adding features and getting the existing ones right. Authentication "
     "and a database would have been good additions, but each one could have eaten a week or two. I made the "
     "call to use browser based storage and keep those items as future work, so I could focus on the core "
     "student journey. Saying no to good ideas, on purpose, was uncomfortable, but it kept the project on "
     "track and it shipped.")

cue("WHAT I AM PROUD OF  (about 6:15 - 7:00)")
para("What I am most proud of is that the tool actually works and feels thoughtful. You paste your resume "
     "and get real feedback in seconds. You explore a career path and the suggestions fit your background. "
     "You practice an interview and get honest coaching. And it all connects.")
para("This project also taught me something about myself. I care most about the part of the product where a "
     "real person feels value. I spent more time on the prompts and the user experience than on the backend, "
     "because that is where a confused student turns into a prepared one. I also learned that I can manage "
     "scope. I resisted the urge to add everything, and I stayed focused on the journey.")

cue("FUTURE AND CLOSING  (about 7:00 - 7:45)")
para("If I keep building, the next steps are clear. I would add real accounts and a database so students can "
     "save their history and track progress over time. I would test different prompts with real users instead "
     "of relying on my own judgment. And I would make the mobile experience just as strong as the desktop one.")
para("If I had to sum up this capstone in one sentence, it is this. The magic of AI is not in the model. It is "
     "in thoughtful integration and sharp prompt design. Claude is powerful, but the real work was "
     "understanding what students need and building an experience that feels natural. Pathway AI is live, it "
     "is free, and it takes students from confused to career ready. Thank you for watching.")

out = os.path.join(HERE, "Pathway_AI_Video_Reflection_Script.docx")
doc.save(out)
print("wrote", out)
